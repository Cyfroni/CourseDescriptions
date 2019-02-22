import re
from collections import defaultdict
from time import sleep

import requests
from docx import Document
from googletrans import Translator

SOURCE = "https://www.one-tab.com/page/NKYb3JWuQMmmIfZRy2A91g"
HTTP = 'https://usosweb.usos.pw.edu.pl/'
URL_pattern = re.escape(f"""href="{HTTP}""") + "(.*?)" + re.escape('">')
TABLE_pattern = re.escape("<table class='grey' cellspacing='1px'>") + "(.*?)" + re.escape('</table>')
translator = Translator()


def get_utf(_source):
    return requests.get(_source).content.decode("utf-8")


def get_rows(_table):
    return re.findall('<tr.*?>(.*?)</tr>', _table, re.S)


def get_columns(_row):
    return re.findall('<td.*?>(.*?)</td>', _row, re.S)


def clean_html(_raw_html):
    _raw_html = _raw_html.replace(r'<br>', '\n')
    _raw_text = re.sub(r'<.*?>', '', _raw_html)
    # _raw_text = unicodedata.normalize('NFD', _raw_text).encode('utf-8', 'ignore')
    # _raw_text = re.sub(r'[^\x00-\x7F]+', '', _raw_text)
    _raw_text = re.sub(r'^\s*', '', _raw_text)
    _raw_text = re.sub(r'\s{2,}', ' ', _raw_text)
    return _raw_text  # if _raw_text else "None"


def description_txt(_data):
    return f"""\
Name:               {_data['Nazwa przedmiotu:']}
Name in English:    {_data['Nazwa przedmiotu:_en']}
ECTS points:        {_data['Punkty ECTS i inne:']}

### Short description ###
{_data['Skrócony opis:_en']}

@@@ Skrócony opis @@@
{_data['Skrócony opis:']}

### Description ###
{_data['Pełny opis:_en']}

@@@ Pełny opis @@@
{_data['Pełny opis:']}

### Bibliography ###
{_data['Literatura:_en']}

@@@ Literatura @@@
{_data['Literatura:']}

### Assessment methods and assessment criteria ###
{_data['Metody i kryteria oceniania:_en']}

@@@ Metody i kryteria oceniania @@@
{_data['Metody i kryteria oceniania:']}\
"""


def add_paragraph(_document, _data, _translation, _elem):
    polish = _elem[0]
    english = _elem[1]
    _document.add_heading(english, level=1)
    _document.add_paragraph(_data[polish + '_en'])
    if _translation:
        _document.add_paragraph().add_run(_data[polish]).italic = True


def description_docx(_data, _translation=True):
    name = _data['Nazwa przedmiotu:_en']
    file_name = f"[{_data['Kod wydziałowy:']}]{name}.docx"

    document = Document()
    document.add_heading(name, 0)

    lines = [
        ['Punkty ECTS i inne:', 'ECTS points'],
        ['Skrócony opis:', 'Short description'],
        ['Pełny opis:', 'Description'],
        ['Literatura:', 'Bibliography'],
        ['Metody i kryteria oceniania:', 'Assessment methods and assessment criteria'],
    ]

    for line in lines:
        add_paragraph(document, _data, _translation, line)

    document.save(file_name)


def write_to_file(_data, log=True):
    file_name = f"[{_data['Kod wydziałowy:']}]{_data['Nazwa przedmiotu:_en']}.txt"
    with open(file_name, 'w+') as file:
        file.truncate(0)
        content = description_txt(_data)
        if log:
            print(content)
        file.write(content)


def translate(_text):
    sleep(10)
    return translator.translate(text).text


translator.translate('안녕하세요.')  # test

urls = re.findall(URL_pattern, get_utf(SOURCE))

for rest in urls:
    data = defaultdict(lambda: r"BRAK !!!")
    url = (HTTP + rest).replace('amp;', '')
    table = re.findall(TABLE_pattern, get_utf(url), re.S)[0]

    for row in get_rows(table):
        cols = get_columns(row)
        text = clean_html(cols[1])

        data[cols[0]] = text
        data[cols[0] + '_en'] = translate(text)

    write_to_file(data)
    description_docx(data)
